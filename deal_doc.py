import os

import docx
import jieba
import pandas as pd

#工具类，包含各种处理txt，csv和doc的方法
class DealtextUtil(object):
    # def __init__(self):

    def deal_text(self, text):
        a = text.find('如不服本判决')
        text = text[:a]
        return text

    #按句子分割
    def readdoc(self,path):
        doc1 = docx.Document(path)
        textc = ''
        for i in range(len(doc1.paragraphs)):
            text1 = doc1.paragraphs[i].text
            textc += text1
        textc = self.deal_text(textc)
        textlist=textc.split('。')
        res=[]
        for textc in textlist:
            textwb = jieba.cut(textc, cut_all=False)
            str_out = ' '.join(textwb).replace('，', '').replace('。', '').replace('？', '').replace('！', '') \
                .replace('“', '').replace('”', '').replace('：', '').replace('…', '').replace('（', '').replace('）', '') \
                .replace('—', '').replace('《', '').replace('》', '').replace('、', '').replace('‘', '') \
                .replace('’', '')
            list = str_out.split(' ')
            while '' in list:
                list.remove('')
            res.append(list)
        return res

    #按词分割：
    def readdoc_word(self,path):
        doc1 = docx.Document(path)
        textc = ''
        for i in range(len(doc1.paragraphs)):
            text1 = doc1.paragraphs[i].text
            textc += text1
        textc = self.deal_text(textc)
        textwb = jieba.cut(textc, cut_all=False)
        str_out = ' '.join(textwb).replace('，', '').replace('。', '').replace('？', '').replace('！', '') \
            .replace('“', '').replace('”', '').replace('：', '').replace('…', '').replace('（', '').replace('）', '') \
            .replace('—', '').replace('《', '').replace('》', '').replace('、', '').replace('‘', '') \
            .replace('’', '')
        list = str_out.split(' ')
        key_index=list.index('本院认为')
        start=max(0,key_index-200)
        end=min(key_index+100,len(list)-1)
        list=list[start:end]
        return list

    def readtext(self,path):
        doc1 = docx.Document(path)
        textc = ''
        for i in range(len(doc1.paragraphs)):
            text1 = doc1.paragraphs[i].text
            textc += text1
        textc = self.deal_text(textc)
        return textc

    #得到分词后的文本list
    def get_doclist(self,path, file):
        filename = os.path.splitext(file)[0]
        doc1 = docx.Document(os.path.join(path, file))

        doc_fenci = self.readdoc(doc1)
        return doc_fenci, filename

    #以dict的形式返回每个目录下的文件数{文件名：文件数}
    def get_filecount(paths):
        count = {}
        for path in paths:
            count[path] = len(os.listdir(path))
        return count

    #读取csv文件
    def readdata(self, path):
        data = pd.read_csv(path, encoding='gbk')
        return data




